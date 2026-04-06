from __future__ import annotations

"""
Microbiology Lecture Notes Builder
=================================
Turn lecture videos/audio into student-friendly English Word notes.

Core pipeline
-------------
1) Extract audio from video.
2) Transcribe speech with faster-whisper (if installed).
3) Sample video frames and extract slide/screen text with OCR.
4) Merge spoken and visual content into logical chunks.
5) Rewrite into structured study notes using an LLM provider.
6) Export polished DOCX and JSON outputs.
"""

from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
import json
import os
import re
import shutil
import subprocess
import tempfile
from collections import defaultdict


@dataclass
class TranscriptSegment:
    start_sec: float
    end_sec: float
    text: str
    speaker: Optional[str] = None


@dataclass
class SlideText:
    timestamp_sec: float
    frame_path: str
    extracted_text: str


@dataclass
class LectureChunk:
    title: str
    raw_transcript: str
    slide_text: str
    merged_context: str
    start_sec: float = 0.0
    end_sec: float = 0.0


@dataclass
class NoteSection:
    title: str
    learning_goal: str
    key_points: List[str]
    comparison_table: List[Dict[str, str]] = field(default_factory=list)
    common_mistakes: List[str] = field(default_factory=list)
    quick_recall: List[str] = field(default_factory=list)
    mini_mind_map: List[str] = field(default_factory=list)


@dataclass
class StudyNotes:
    lecture_title: str
    language: str
    audience_level: str
    summary: List[str]
    sections: List[NoteSection]
    glossary: List[Dict[str, str]]
    exam_style_questions: List[Dict[str, Any]]


@dataclass
class ToolConfig:
    output_language: str = "English"
    audience_level: str = "Medical foundation students"
    max_chunk_chars: int = 4200
    frame_interval_sec: int = 20
    enable_ocr: bool = True
    llm_model_name: str = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
    transcription_model_name: str = "base"
    whisper_device: str = "cpu"
    whisper_compute_type: str = "int8"
    ocr_lang: str = "eng"
    min_ocr_text_length: int = 12
    keep_temp_files: bool = False


def run_command(cmd: List[str]) -> None:
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"Command failed: {' '.join(cmd)}\n{result.stderr}")


def slugify(text: str) -> str:
    text = text.strip().lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-") or "output"


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def dedupe_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for item in items:
        key = normalize_whitespace(item).lower()
        if key and key not in seen:
            seen.add(key)
            out.append(normalize_whitespace(item))
    return out


def clean_ocr_text(text: str) -> str:
    text = text.replace("|", "I")
    text = re.sub(r"[_=]{2,}", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_sentences(text: str) -> List[str]:
    text = normalize_whitespace(text)
    if not text:
        return []
    return [p.strip() for p in re.split(r"(?<=[.!?])\s+", text) if p.strip()]


def ensure_ffmpeg() -> None:
    if shutil.which("ffmpeg") is None:
        raise EnvironmentError("ffmpeg is not installed or not available in PATH.")


class MediaProcessor:
    def extract_audio(self, media_path: Path, out_wav: Path) -> Path:
        ensure_ffmpeg()
        run_command([
            "ffmpeg", "-y", "-i", str(media_path),
            "-ac", "1", "-ar", "16000", str(out_wav)
        ])
        return out_wav

    def extract_frames(self, media_path: Path, out_dir: Path, interval_sec: int) -> List[Path]:
        ensure_ffmpeg()
        out_dir.mkdir(parents=True, exist_ok=True)
        pattern = out_dir / "frame_%05d.jpg"
        run_command([
            "ffmpeg", "-y", "-i", str(media_path),
            "-vf", f"fps=1/{interval_sec}",
            str(pattern)
        ])
        return sorted(out_dir.glob("frame_*.jpg"))


class TranscriptionProvider:
    def transcribe(self, audio_path: Path) -> List[TranscriptSegment]:
        raise NotImplementedError


class OCRProvider:
    def extract_text(self, image_path: Path) -> str:
        raise NotImplementedError


class LLMProvider:
    def build_notes(self, lecture_title: str, chunks: List[LectureChunk], config: ToolConfig) -> StudyNotes:
        raise NotImplementedError


class FasterWhisperTranscriptionProvider(TranscriptionProvider):
    def __init__(self, model_name: str = "base", device: str = "cpu", compute_type: str = "int8"):
        self.model_name = model_name
        self.device = device
        self.compute_type = compute_type

    def transcribe(self, audio_path: Path) -> List[TranscriptSegment]:
        from faster_whisper import WhisperModel

        model = WhisperModel(self.model_name, device=self.device, compute_type=self.compute_type)
        segments, _info = model.transcribe(str(audio_path), beam_size=5, vad_filter=True)

        output: List[TranscriptSegment] = []
        for seg in segments:
            text = normalize_whitespace(seg.text)
            if text:
                output.append(TranscriptSegment(
                    start_sec=float(seg.start),
                    end_sec=float(seg.end),
                    text=text,
                ))
        return output


class DummyTranscriptionProvider(TranscriptionProvider):
    def transcribe(self, audio_path: Path) -> List[TranscriptSegment]:
        sample_text = (
            "Introduction to virology. Definition of viruses. Differences between viruses and bacteria. "
            "Historical examples include smallpox, rabies, tobacco mosaic virus, and bacteriophages."
        )
        return [TranscriptSegment(0, 60, sample_text)]


class TesseractOCRProvider(OCRProvider):
    def __init__(self, lang: str = "eng"):
        self.lang = lang

    def extract_text(self, image_path: Path) -> str:
        import pytesseract
        from PIL import Image, ImageFilter, ImageOps

        image = Image.open(image_path).convert("L")
        image = ImageOps.autocontrast(image)
        image = image.filter(ImageFilter.SHARPEN)
        text = pytesseract.image_to_string(image, lang=self.lang, config="--psm 6")
        return clean_ocr_text(text)


class OpenAILLMProvider(LLMProvider):
    def __init__(self, model_name: str):
        self.model_name = model_name

    def build_notes(self, lecture_title: str, chunks: List[LectureChunk], config: ToolConfig) -> StudyNotes:
        from openai import OpenAI

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY is not set.")

        client = OpenAI(api_key=api_key)
        prompt = self._make_prompt(lecture_title, chunks, config)
        response = client.responses.create(
            model=self.model_name,
            input=prompt,
            temperature=0.2,
        )
        data = json.loads(self._extract_json(response.output_text.strip()))
        return self._parse_study_notes(data)

    def _make_prompt(self, lecture_title: str, chunks: List[LectureChunk], config: ToolConfig) -> str:
        chunks_json = json.dumps([
            {
                "title": c.title,
                "start": c.start_sec,
                "end": c.end_sec,
                "transcript": c.raw_transcript,
                "slide_text": c.slide_text,
                "merged_context": c.merged_context,
            }
            for c in chunks
        ], ensure_ascii=False, indent=2)

        return REAL_LLM_PROMPT_TEMPLATE.format(
            lecture_title=lecture_title,
            audience_level=config.audience_level,
            output_language=config.output_language,
            chunks_json=chunks_json,
        )

    def _extract_json(self, text: str) -> str:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            raise ValueError("LLM output did not contain valid JSON object.")
        return match.group(0)

    def _parse_study_notes(self, data: Dict[str, Any]) -> StudyNotes:
        sections = [
            NoteSection(
                title=s.get("title", "Untitled Section"),
                learning_goal=s.get("learning_goal", ""),
                key_points=s.get("key_points", []) or [],
                comparison_table=s.get("comparison_table", []) or [],
                common_mistakes=s.get("common_mistakes", []) or [],
                quick_recall=s.get("quick_recall", []) or [],
                mini_mind_map=s.get("mini_mind_map", []) or [],
            )
            for s in data.get("sections", [])
        ]
        return StudyNotes(
            lecture_title=data.get("lecture_title", "Lecture Notes"),
            language=data.get("language", "English"),
            audience_level=data.get("audience_level", "Medical foundation students"),
            summary=data.get("summary", []) or [],
            sections=sections,
            glossary=data.get("glossary", []) or [],
            exam_style_questions=data.get("exam_style_questions", []) or [],
        )


class RuleBasedLLMProvider(LLMProvider):
    def build_notes(self, lecture_title: str, chunks: List[LectureChunk], config: ToolConfig) -> StudyNotes:
        full_text = "\n\n".join(c.merged_context for c in chunks)
        raw_sentences = split_sentences(full_text)
        key_points = dedupe_preserve_order(self._prioritize_sentences(raw_sentences))[:12]
        sections = self._build_sections_from_chunks(chunks)
        glossary = self._build_glossary(full_text)
        questions = self._build_exam_questions(full_text)
        summary = dedupe_preserve_order([s for s in key_points[:5] if 25 <= len(s) <= 180])

        if not sections:
            sections = [
                NoteSection(
                    title="Core Concepts",
                    learning_goal="Understand the main lecture ideas in a simple study format.",
                    key_points=key_points[:8] or ["No key points extracted."],
                    comparison_table=self._maybe_build_comparison_table(full_text),
                    common_mistakes=self._default_common_mistakes(full_text),
                    quick_recall=self._quick_recall_from_text(full_text),
                    mini_mind_map=self._mind_map_from_text(full_text),
                )
            ]

        return StudyNotes(
            lecture_title=lecture_title,
            language=config.output_language,
            audience_level=config.audience_level,
            summary=summary or ["Structured notes were generated from the lecture content."],
            sections=sections,
            glossary=glossary,
            exam_style_questions=questions,
        )

    def _build_sections_from_chunks(self, chunks: List[LectureChunk]) -> List[NoteSection]:
        sections: List[NoteSection] = []
        for idx, chunk in enumerate(chunks, start=1):
            text = chunk.merged_context
            title = self._infer_section_title(text, idx)
            key_points = dedupe_preserve_order(self._prioritize_sentences(split_sentences(text)))[:8]
            if not key_points:
                continue
            sections.append(NoteSection(
                title=title,
                learning_goal=self._infer_learning_goal(title),
                key_points=key_points,
                comparison_table=self._maybe_build_comparison_table(text),
                common_mistakes=self._default_common_mistakes(text),
                quick_recall=self._quick_recall_from_text(text),
                mini_mind_map=self._mind_map_from_text(text),
            ))
        return sections[:6]

    def _infer_section_title(self, text: str, idx: int) -> str:
        lowered = text.lower()
        candidates = [
            ("Definition and Nature of Viruses", ["definition", "virus", "capsid", "envelope"]),
            ("Historical Development of Virology", ["history", "jenner", "pasteur", "smallpox", "rabies"]),
            ("Virus Structure and Replication", ["genome", "capsid", "intracellular", "replicate", "virion"]),
            ("Virus vs Bacteria", ["bacteria", "binary fission", "antibiotics", "metabolically"]),
            ("Useful and Harmful Effects of Viruses", ["useful", "pesticide", "anticancer", "pandemic", "zoonotic"]),
            ("Atypical Infectious Agents", ["viroid", "prion", "satellite", "virusoid"]),
        ]
        best_title = f"Section {idx}"
        best_score = 0
        for title, keywords in candidates:
            score = sum(1 for kw in keywords if kw in lowered)
            if score > best_score:
                best_title = title
                best_score = score
        return best_title

    def _infer_learning_goal(self, title: str) -> str:
        mapping = {
            "Definition and Nature of Viruses": "Define viruses and explain what makes them different from cellular organisms.",
            "Historical Development of Virology": "Follow the major discoveries that built the science of virology.",
            "Virus Structure and Replication": "Explain the main structural parts of a virus and how replication depends on host cells.",
            "Virus vs Bacteria": "Differentiate between viral and bacterial properties in a way that supports exam recall.",
            "Useful and Harmful Effects of Viruses": "Connect viral disease burden with beneficial scientific and medical uses.",
            "Atypical Infectious Agents": "Compare viruses with viroids, satellites, virusoids, and prions.",
        }
        return mapping.get(title, "Understand and memorize the main ideas in this part of the lecture.")

    def _prioritize_sentences(self, sentences: List[str]) -> List[str]:
        scored: List[Tuple[int, str]] = []
        signal_words = {
            "define": 4, "definition": 4, "important": 3, "means": 2,
            "virus": 2, "viruses": 2, "history": 2, "replicate": 2,
            "difference": 2, "compare": 2, "cause": 2, "used": 2,
            "example": 1, "examples": 1, "because": 1,
        }
        for s in sentences:
            score = 0
            low = s.lower()
            for word, weight in signal_words.items():
                if word in low:
                    score += weight
            if 35 <= len(s) <= 220:
                score += 2
            scored.append((score, s))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [s for _, s in scored]

    def _maybe_build_comparison_table(self, text: str) -> List[Dict[str, str]]:
        lowered = text.lower()
        if "bacteria" in lowered and "virus" in lowered:
            return [
                {"Aspect": "Cellularity", "Virus": "Acellular", "Bacteria": "Cellular unicellular organism"},
                {"Aspect": "Genetic material", "Virus": "DNA or RNA only", "Bacteria": "Contain both DNA and RNA"},
                {"Aspect": "Replication", "Virus": "Inside host cell only", "Bacteria": "Binary fission"},
                {"Aspect": "Ribosomes", "Virus": "Absent", "Bacteria": "Present"},
                {"Aspect": "Response to antibiotics", "Virus": "Not treated by antibiotics", "Bacteria": "Often sensitive to antibiotics"},
            ]
        if "viroid" in lowered or "prion" in lowered or "satellite" in lowered:
            return [
                {"Agent": "Viroid", "Core feature": "Small circular RNA without protein coat"},
                {"Agent": "Satellite virus", "Core feature": "Needs helper virus for replication"},
                {"Agent": "Virusoid", "Core feature": "Small circular RNA using helper virus coat"},
                {"Agent": "Prion", "Core feature": "Infectious misfolded protein without nucleic acid"},
            ]
        return []

    def _default_common_mistakes(self, text: str) -> List[str]:
        hints = []
        lowered = text.lower()
        if "dna or rna" in lowered or ("dna" in lowered and "rna" in lowered):
            hints.append("Do not say that a virus contains both DNA and RNA. It usually contains one type only.")
        if "antibiotics" in lowered:
            hints.append("Do not use antibiotics as a treatment for viral infection unless there is a bacterial co-infection.")
        if "prion" in lowered:
            hints.append("Do not classify prions as viruses. Prions are infectious proteins, not viral particles.")
        if "satellite" in lowered or "virusoid" in lowered:
            hints.append("Do not confuse satellite virus with virusoid. The satellite virus encodes its own coat, while virusoid uses the helper virus coat.")
        if not hints:
            hints = [
                "Avoid copying the lecture wording without understanding the concept.",
                "Always connect structure with function when reviewing microbiology topics.",
            ]
        return dedupe_preserve_order(hints)[:4]

    def _quick_recall_from_text(self, text: str) -> List[str]:
        facts = []
        lowered = text.lower()
        if "virion" in lowered:
            facts.append("Virion = complete mature infectious viral particle.")
        if "capsid" in lowered:
            facts.append("Capsid = protein coat that protects viral nucleic acid.")
        if "envelope" in lowered:
            facts.append("Envelope = lipid layer surrounding the capsid in enveloped viruses.")
        if "intracellular" in lowered:
            facts.append("Viruses are obligate intracellular agents.")
        if "smallpox" in lowered:
            facts.append("Smallpox was declared eradicated by WHO in 1980 in the lecture content.")
        if "rabies" in lowered:
            facts.append("Rabies is classically fatal once symptoms are established.")
        return dedupe_preserve_order(facts)[:6]

    def _mind_map_from_text(self, text: str) -> List[str]:
        lowered = text.lower()
        lines = []
        if "history" in lowered or "smallpox" in lowered or "jenner" in lowered:
            lines.append("Virology -> History -> Smallpox -> Variolation -> Vaccination -> Jenner")
        if "virus" in lowered and "bacteria" in lowered:
            lines.append("Microbes -> Virus vs Bacteria -> Structure -> Replication -> Treatment")
        if "prion" in lowered or "viroid" in lowered:
            lines.append("Infectious agents -> Viruses -> Atypical agents -> Viroids / Satellites / Prions")
        if not lines:
            lines.append("Lecture -> Core concept -> Key features -> Clinical relevance -> Exam recall")
        return lines[:4]

    def _build_glossary(self, text: str) -> List[Dict[str, str]]:
        glossary_map = {
            "virion": "Complete mature infectious viral particle.",
            "capsid": "Protein coat surrounding the viral genome.",
            "envelope": "Outer lipid membrane present in some viruses.",
            "viroid": "Small circular RNA infectious agent without a protein coat.",
            "prion": "Infectious misfolded protein lacking nucleic acid.",
            "bacteriophage": "Virus that infects bacteria.",
            "attenuated vaccine": "Vaccine made from a weakened form of the pathogen.",
            "variolation": "Early protective method using material from smallpox lesions.",
        }
        lowered = text.lower()
        return [{"term": term.title(), "definition": definition} for term, definition in glossary_map.items() if term in lowered][:10]

    def _build_exam_questions(self, text: str) -> List[Dict[str, Any]]:
        questions = [{
            "type": "MCQ",
            "question": "Which statement best describes a virus?",
            "choices": [
                "A living cell with ribosomes",
                "An acellular agent that replicates only inside host cells",
                "A bacterium lacking a cell wall",
                "A fungus with RNA only",
            ],
            "answer": "An acellular agent that replicates only inside host cells",
        }]
        lowered = text.lower()
        if "prion" in lowered:
            questions.append({
                "type": "MCQ",
                "question": "Which infectious agent lacks nucleic acid?",
                "choices": ["Virion", "Viroid", "Prion", "Satellite virus"],
                "answer": "Prion",
            })
        if "bacteria" in lowered:
            questions.append({
                "type": "Short Answer",
                "question": "Mention two major differences between viruses and bacteria.",
                "answer": "Examples: viruses are acellular and replicate only inside host cells, while bacteria are cellular and reproduce by binary fission.",
            })
        return questions[:6]


class LectureNotesBuilder:
    def __init__(
        self,
        media_processor: MediaProcessor,
        transcription_provider: TranscriptionProvider,
        ocr_provider: Optional[OCRProvider],
        llm_provider: LLMProvider,
        config: ToolConfig,
    ):
        self.media_processor = media_processor
        self.transcription_provider = transcription_provider
        self.ocr_provider = ocr_provider
        self.llm_provider = llm_provider
        self.config = config

    def process(self, media_path: str, lecture_title: Optional[str] = None) -> StudyNotes:
        media = Path(media_path)
        if not media.exists():
            raise FileNotFoundError(f"Input not found: {media}")

        title = lecture_title or media.stem.replace("_", " ").title()
        if self.config.keep_temp_files:
            temp_root = Path(tempfile.mkdtemp(prefix="microbio_notes_"))
            notes = self._run_pipeline(media, title, temp_root)
            print(f"Temporary files kept at: {temp_root}")
            return notes

        with tempfile.TemporaryDirectory(prefix="microbio_notes_") as tmp:
            return self._run_pipeline(media, title, Path(tmp))

    def _run_pipeline(self, media: Path, title: str, tmp_dir: Path) -> StudyNotes:
        wav_path = tmp_dir / "audio.wav"
        frames_dir = tmp_dir / "frames"

        self.media_processor.extract_audio(media, wav_path)
        transcript_segments = self.transcription_provider.transcribe(wav_path)

        slide_texts: List[SlideText] = []
        if self.config.enable_ocr and self.ocr_provider:
            frame_paths = self.media_processor.extract_frames(media, frames_dir, self.config.frame_interval_sec)
            for idx, frame_path in enumerate(frame_paths):
                try:
                    text = self.ocr_provider.extract_text(frame_path)
                except Exception:
                    text = ""
                text = normalize_whitespace(text)
                if len(text) >= self.config.min_ocr_text_length:
                    slide_texts.append(SlideText(
                        timestamp_sec=idx * self.config.frame_interval_sec,
                        frame_path=str(frame_path),
                        extracted_text=text,
                    ))

        chunks = self._merge_into_chunks(transcript_segments, slide_texts)
        return self.llm_provider.build_notes(title, chunks, self.config)

    def _merge_into_chunks(
        self,
        transcript_segments: List[TranscriptSegment],
        slide_texts: List[SlideText],
    ) -> List[LectureChunk]:
        if not transcript_segments:
            transcript_segments = [TranscriptSegment(0, 0, "")]

        slide_index: Dict[int, List[str]] = defaultdict(list)
        window = max(self.config.frame_interval_sec, 1)
        for item in slide_texts:
            bucket = int(item.timestamp_sec // window)
            slide_index[bucket].append(item.extracted_text)

        assembled_units: List[Tuple[float, float, str, str, str]] = []
        for seg in transcript_segments:
            bucket = int(seg.start_sec // window)
            slide_text = "\n".join(dedupe_preserve_order(slide_index.get(bucket, [])))
            merged_context = (
                f"Transcript:\n{seg.text}\n\n"
                f"Visual text near this moment:\n{slide_text}"
            ).strip()
            assembled_units.append((seg.start_sec, seg.end_sec, seg.text, slide_text, merged_context))

        chunks: List[LectureChunk] = []
        current_transcript: List[str] = []
        current_slides: List[str] = []
        current_merged: List[str] = []
        current_start = assembled_units[0][0]
        current_end = assembled_units[0][1]
        current_len = 0

        for start, end, transcript, slide_text, merged in assembled_units:
            estimated_add = len(merged) + 2
            if current_len + estimated_add > self.config.max_chunk_chars and current_merged:
                chunks.append(LectureChunk(
                    title=f"Chunk {len(chunks) + 1}",
                    raw_transcript="\n".join(current_transcript).strip(),
                    slide_text="\n".join(dedupe_preserve_order(current_slides)).strip(),
                    merged_context="\n\n".join(current_merged).strip(),
                    start_sec=current_start,
                    end_sec=current_end,
                ))
                current_transcript = [transcript]
                current_slides = [slide_text] if slide_text else []
                current_merged = [merged]
                current_start = start
                current_end = end
                current_len = len(merged)
            else:
                current_transcript.append(transcript)
                if slide_text:
                    current_slides.append(slide_text)
                current_merged.append(merged)
                current_end = end
                current_len += estimated_add

        if current_merged:
            chunks.append(LectureChunk(
                title=f"Chunk {len(chunks) + 1}",
                raw_transcript="\n".join(current_transcript).strip(),
                slide_text="\n".join(dedupe_preserve_order(current_slides)).strip(),
                merged_context="\n\n".join(current_merged).strip(),
                start_sec=current_start,
                end_sec=current_end,
            ))
        return chunks


REAL_LLM_PROMPT_TEMPLATE = """
You are building microbiology-style study notes for {audience_level}.
Lecture title: {lecture_title}
Output language: {output_language}

Strict rules:
- Do NOT copy the transcript word-for-word.
- Rewrite for understanding and memorization.
- Keep the science accurate.
- Use short, high-yield bullet points.
- Group overlapping ideas.
- Include comparison tables when useful.
- Include mini mind maps as text trees or arrow chains.
- Add common mistakes students may make.
- Add glossary terms.
- Add exam-style questions.
- Prefer clear undergraduate-level English.
- Avoid unnecessary jargon.

Input chunks:
{chunks_json}

Return valid JSON only with this structure:
{{
  "lecture_title": "...",
  "language": "...",
  "audience_level": "...",
  "summary": ["..."],
  "sections": [
    {{
      "title": "...",
      "learning_goal": "...",
      "key_points": ["..."],
      "comparison_table": [{{"Aspect": "...", "Feature": "..."}}],
      "common_mistakes": ["..."],
      "quick_recall": ["..."],
      "mini_mind_map": ["..."]
    }}
  ],
  "glossary": [{{"term": "...", "definition": "..."}}],
  "exam_style_questions": [
    {{"type": "MCQ", "question": "...", "choices": ["..."], "answer": "..."}}
  ]
}}
""".strip()


def _style_doc(document) -> None:
    from docx.shared import Pt
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def export_to_docx(notes: StudyNotes, output_path: str) -> str:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _style_doc(doc)

    title = doc.add_heading(notes.lecture_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Language: {notes.language}    ").italic = True
    meta.add_run(f"Audience: {notes.audience_level}").italic = True

    doc.add_heading("High-Yield Summary", level=1)
    for item in notes.summary:
        doc.add_paragraph(item, style="List Bullet")

    for section in notes.sections:
        doc.add_heading(section.title, level=1)
        p = doc.add_paragraph()
        p.add_run("Learning goal: ").bold = True
        p.add_run(section.learning_goal)

        doc.add_heading("Key Points", level=2)
        for point in section.key_points:
            doc.add_paragraph(point, style="List Bullet")

        if section.comparison_table:
            doc.add_heading("Comparison Table", level=2)
            headers = list(section.comparison_table[0].keys())
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            for row in section.comparison_table:
                cells = table.add_row().cells
                for i, header in enumerate(headers):
                    cells[i].text = str(row.get(header, ""))

        if section.common_mistakes:
            doc.add_heading("Common Mistakes", level=2)
            for item in section.common_mistakes:
                doc.add_paragraph(item, style="List Bullet")

        if section.quick_recall:
            doc.add_heading("Quick Recall", level=2)
            for item in section.quick_recall:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item).bold = True

        if section.mini_mind_map:
            doc.add_heading("Mini Mind Map", level=2)
            for item in section.mini_mind_map:
                doc.add_paragraph(item, style="List Bullet")

    if notes.glossary:
        doc.add_heading("Glossary", level=1)
        for entry in notes.glossary:
            para = doc.add_paragraph()
            para.add_run(entry.get("term", "") + ": ").bold = True
            para.add_run(entry.get("definition", ""))

    if notes.exam_style_questions:
        doc.add_heading("Exam-Style Questions", level=1)
        for i, q in enumerate(notes.exam_style_questions, start=1):
            qp = doc.add_paragraph()
            qp.add_run(f"Q{i}. ").bold = True
            qp.add_run(q.get("question", ""))
            for choice in q.get("choices", []):
                doc.add_paragraph(choice, style="List Bullet")
            if q.get("answer"):
                ans = doc.add_paragraph()
                ans.add_run("Answer: ").bold = True
                ans.add_run(str(q["answer"]))

    doc.save(output_path)
    return output_path


def export_to_json(notes: StudyNotes, output_path: str) -> str:
    Path(output_path).write_text(json.dumps(asdict(notes), ensure_ascii=False, indent=2), encoding="utf-8")
    return output_path


def build_transcription_provider(config: ToolConfig, mode: str) -> TranscriptionProvider:
    if mode == "dummy":
        return DummyTranscriptionProvider()
    return FasterWhisperTranscriptionProvider(
        model_name=config.transcription_model_name,
        device=config.whisper_device,
        compute_type=config.whisper_compute_type,
    )


def build_ocr_provider(config: ToolConfig) -> Optional[OCRProvider]:
    if not config.enable_ocr:
        return None
    return TesseractOCRProvider(lang=config.ocr_lang)


def build_llm_provider(config: ToolConfig, mode: str) -> LLMProvider:
    if mode == "rule":
        return RuleBasedLLMProvider()
    return OpenAILLMProvider(model_name=config.llm_model_name)


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(description="Turn lecture media into student-friendly study notes.")
    parser.add_argument("input_media", help="Path to input video/audio file")
    parser.add_argument("--title", default=None, help="Optional lecture title override")
    parser.add_argument("--lang", default="English", help="Output language")
    parser.add_argument("--json-out", default=None, help="Optional path to JSON output")
    parser.add_argument("--docx-out", default=None, help="Optional path to DOCX output")
    parser.add_argument("--frame-interval", type=int, default=20, help="Frame extraction interval in seconds")
    parser.add_argument("--max-chars", type=int, default=4200, help="Max chars per merged chunk")
    parser.add_argument("--transcriber", choices=["auto", "dummy"], default="auto")
    parser.add_argument("--writer", choices=["auto", "rule"], default="auto")
    parser.add_argument("--disable-ocr", action="store_true")
    parser.add_argument("--keep-temp", action="store_true")
    args = parser.parse_args()

    config = ToolConfig(
        output_language=args.lang,
        frame_interval_sec=args.frame_interval,
        max_chunk_chars=args.max_chars,
        enable_ocr=not args.disable_ocr,
        keep_temp_files=args.keep_temp,
    )

    if args.transcriber == "auto":
        try:
            transcription_provider = build_transcription_provider(config, "auto")
        except Exception:
            transcription_provider = DummyTranscriptionProvider()
            print("[warn] faster-whisper unavailable. Falling back to dummy transcription.")
    else:
        transcription_provider = DummyTranscriptionProvider()

    try:
        ocr_provider = build_ocr_provider(config)
    except Exception:
        ocr_provider = None
        print("[warn] OCR provider unavailable. Continuing without OCR.")

    if args.writer == "auto":
        try:
            llm_provider = build_llm_provider(config, "auto")
            if not os.getenv("OPENAI_API_KEY"):
                raise RuntimeError("Missing OPENAI_API_KEY")
        except Exception:
            llm_provider = RuleBasedLLMProvider()
            print("[warn] OpenAI writer unavailable. Falling back to rule-based notes builder.")
    else:
        llm_provider = RuleBasedLLMProvider()

    builder = LectureNotesBuilder(
        media_processor=MediaProcessor(),
        transcription_provider=transcription_provider,
        ocr_provider=ocr_provider,
        llm_provider=llm_provider,
        config=config,
    )

    notes = builder.process(args.input_media, lecture_title=args.title)
    base_name = slugify(notes.lecture_title)
    json_out = args.json_out or f"{base_name}.json"
    docx_out = args.docx_out or f"{base_name}.docx"

    export_to_json(notes, json_out)
    export_to_docx(notes, docx_out)

    print(f"JSON saved to: {json_out}")
    print(f"DOCX saved to: {docx_out}")


if __name__ == "__main__":
    main()